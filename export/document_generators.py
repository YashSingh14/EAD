import io
import textwrap
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

def generate_docx_buffer(content: str) -> io.BytesIO:
    """
    1. In-Memory DOCX Generation: Dynamically constructs a Word document 
    and saves it securely into a transient io.BytesIO() binary buffer.
    Zero files are written to the server's local disk.
    """
    doc = Document()
    doc.add_heading('AI Semantic Insights Report', 0)
    
    # Basic structural parsing
    paragraphs = content.split("\n\n")
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0) # Reset pointer to beginning so Streamlit reads from start
    return buffer

def generate_pdf_buffer(content: str) -> io.BytesIO:
    """
    2. In-Memory PDF Generation: Renders a PDF directly to an io.BytesIO() buffer
    using ReportLab's Canvas with automated text wrapping, line breaks, and pagination.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Clean typography configuration
    c.setFont("Helvetica-Bold", 16)
    c.drawString(inch, height - inch, "AI Semantic Insights Report")
    
    c.setFont("Helvetica", 11)
    
    # Begin text drawing routine
    text_object = c.beginText(inch, height - 1.5 * inch)
    
    # Calculate approx maximum characters per line for standard 11pt Helvetica 
    # to prevent text bleeding off the right margin.
    max_chars_per_line = 95
    
    # Split by actual line breaks to respect formatting
    lines = content.split("\n")
    
    for line in lines:
        if not line.strip():
            # Respect empty paragraph spacing
            text_object.textLine("")
            continue
            
        # Wrap long continuous text
        wrapped_lines = textwrap.wrap(line, width=max_chars_per_line)
        for wrapped_line in wrapped_lines:
            # 2. Dynamic Pagination: Check if we are approaching bottom margin
            if text_object.getY() < inch:
                c.drawText(text_object)
                c.showPage() # Create a new page
                c.setFont("Helvetica", 11)
                text_object = c.beginText(inch, height - inch) # Reset cursor to top
            
            # Print line and advance Y cursor automatically
            text_object.textLine(wrapped_line)
            
    # Finalize and push to buffer
    c.drawText(text_object)
    c.save()
    buffer.seek(0)
    
    return buffer
