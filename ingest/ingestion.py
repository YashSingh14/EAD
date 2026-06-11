import os
import logging
import re
from typing import List

import pdfplumber
import docx
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Configure basic logging for the pipeline
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFPlumberLoader:
    """
    Custom LangChain-compatible document loader using pdfplumber for high-fidelity 
    text extraction from complex PDF layouts, multi-column formats, and tables.
    """
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        documents = []
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Extract text; pdfplumber handles reading order naturally much better than PyPDF2
                    text = page.extract_text() or ""
                    
                    # Optionally append table content cleanly if layout requires it
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_text = "\n".join([" | ".join([str(cell).replace('\n', ' ') if cell else "" for cell in row]) for row in table])
                            text += f"\n\n{table_text}\n"

                    if text.strip():
                        metadata = {
                            "source": self.file_path,
                            "filename": os.path.basename(self.file_path),
                            "page_number": i + 1
                        }
                        documents.append(Document(page_content=text.strip(), metadata=metadata))
        except Exception as e:
            logger.error(f"Error reading PDF {self.file_path} with pdfplumber: {e}")
            raise e
            
        return documents

class DocxLoader:
    """
    Custom loader using python-docx (from requirements) to avoid unstructured dependency.
    """
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        documents = []
        try:
            doc = docx.Document(self.file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if text.strip():
                metadata = {
                    "source": self.file_path,
                    "filename": os.path.basename(self.file_path),
                    "page_number": 1 # DOCX lacks easily exposed hard page breaks in pure XML
                }
                documents.append(Document(page_content=text.strip(), metadata=metadata))
        except Exception as e:
            logger.error(f"Error reading DOCX {self.file_path}: {e}")
            raise e
        return documents

class BasicHTMLLoader:
    """
    Lightweight HTML loader using regex to strip tags, keeping dependencies zero-cost.
    """
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        documents = []
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
                
            # Remove styles and scripts
            text = re.sub(r'<style.*?>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<script.*?>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', ' ', text)
            # Normalize whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            if text:
                metadata = {
                    "source": self.file_path,
                    "filename": os.path.basename(self.file_path),
                    "page_number": 1
                }
                documents.append(Document(page_content=text, metadata=metadata))
        except Exception as e:
            logger.error(f"Error reading HTML {self.file_path}: {e}")
            raise e
        return documents

class DocumentIngestionPipeline:
    """
    Robust, multi-format document ingestion and semantic chunking pipeline.
    """
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Optimized semantic chunker aiming for paragraph/sentence boundaries
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""],
            keep_separator=False
        )

    def load_document(self, file_path: str) -> List[Document]:
        """
        Routes the file to the appropriate loader based on extension with fault tolerance.
        """
        ext = os.path.splitext(file_path)[1].lower()
        documents = []

        try:
            if ext == '.pdf':
                logger.info(f"Ingesting PDF: {file_path}")
                loader = PDFPlumberLoader(file_path)
                documents = loader.load()
                
            elif ext == '.docx':
                logger.info(f"Ingesting DOCX: {file_path}")
                loader = DocxLoader(file_path)
                documents = loader.load()
                
            elif ext == '.txt':
                logger.info(f"Ingesting TXT: {file_path}")
                loader = TextLoader(file_path, encoding='utf-8')
                documents = loader.load()
                # Inject filename/page metadata for consistency
                for doc in documents:
                    doc.metadata["filename"] = os.path.basename(file_path)
                    doc.metadata["page_number"] = 1
                    
            elif ext in ['.html', '.htm']:
                logger.info(f"Ingesting HTML: {file_path}")
                loader = BasicHTMLLoader(file_path)
                documents = loader.load()
                
            else:
                logger.warning(f"Skipping unsupported file format ({ext}) for: {file_path}")
                
        except Exception as e:
            logger.error(f"Pipeline skipped malformed or unreadable document {file_path}: {e}")

        return documents

    def process_document(self, file_path: str) -> List[Document]:
        """
        Loads a document and splits it into semantic chunks, preserving metadata.
        """
        documents = self.load_document(file_path)
        if not documents:
            return []
            
        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Successfully created {len(chunks)} chunks for {os.path.basename(file_path)}")
            return chunks
        except Exception as e:
            logger.error(f"Failed to chunk document {file_path}: {e}")
            return []

    def process_directory(self, directory_path: str) -> List[Document]:
        """
        Recursively processes all supported documents in a given directory with complete resilience.
        """
        all_chunks = []
        if not os.path.isdir(directory_path):
            logger.error(f"Directory not found: {directory_path}")
            return all_chunks

        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                chunks = self.process_document(file_path)
                all_chunks.extend(chunks)
                
        logger.info(f"Directory processing complete. Total chunks generated: {len(all_chunks)}")
        return all_chunks

if __name__ == "__main__":
    # Example usage:
    pipeline = DocumentIngestionPipeline(chunk_size=1000, chunk_overlap=200)
    # chunks = pipeline.process_directory("./data/raw_documents")
    # print(f"Sample chunk: {chunks[0].page_content[:100]}... Metadata: {chunks[0].metadata}")
